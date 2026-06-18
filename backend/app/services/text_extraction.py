"""SupportPilot AI — Text Extraction Engine

Extracts text content from uploaded files.
Supports: PDF, DOCX, TXT, Markdown.

Design: Each extractor is a standalone function returning a string.
The TextExtractorFactory selects the right one based on MIME type.
All extractors handle errors gracefully and return meaningful messages.
"""

from __future__ import annotations

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger("supportpilot.extraction")


class BaseExtractor(ABC):
    """Base class for text extractors."""

    @abstractmethod
    def extract(self, file_path: str) -> str:
        """Extract text from a file at the given path."""
        ...

    @abstractmethod
    def extract_from_bytes(self, content: bytes, filename: str = "") -> str:
        """Extract text from a byte buffer."""
        ...


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using PyMuPDF (fitz)."""

    def extract(self, file_path: str) -> str:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return self._fallback_extract(file_path)

        text_parts = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            doc.close()
        except Exception as e:
            logger.error("PDF extraction failed for %s: %s", file_path, e)
            raise ExtractionError(f"Failed to extract PDF: {e}") from e

        return "\n\n".join(text_parts)

    def extract_from_bytes(self, content: bytes, filename: str = "") -> str:
        try:
            import fitz
        except ImportError:
            raise ExtractionError("PyMuPDF not installed. Install with: pip install PyMuPDF")

        text_parts = []
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            doc.close()
        except Exception as e:
            logger.error("PDF extraction from bytes failed: %s", e)
            raise ExtractionError(f"Failed to extract PDF from bytes: {e}") from e

        return "\n\n".join(text_parts)

    def _fallback_extract(self, file_path: str) -> str:
        """Fallback: try pdfplumber if PyMuPDF unavailable."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ExtractionError(
                "No PDF extraction library installed. "
                "Install PyMuPDF: pip install PyMuPDF"
            )


class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files using python-docx."""

    def extract(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError("python-docx not installed. Install with: pip install python-docx")

        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error("DOCX open failed for %s: %s", file_path, e)
            raise ExtractionError(f"Failed to open DOCX: {e}") from e

        return self._extract_from_doc(doc)

    def extract_from_bytes(self, content: bytes, filename: str = "") -> str:
        try:
            from docx import Document
            import io
        except ImportError:
            raise ExtractionError("python-docx not installed. Install with: pip install python-docx")

        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            logger.error("DOCX extraction from bytes failed: %s", e)
            raise ExtractionError(f"Failed to extract DOCX from bytes: {e}") from e

        return self._extract_from_doc(doc)

    def _extract_from_doc(self, doc) -> str:
        """Extract text from a python-docx Document object."""
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        return "\n\n".join(paragraphs)


class TextExtractor(BaseExtractor):
    """Extract text from plain text files (TXT, Markdown)."""

    SUPPORTED_ENCODINGS = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"]

    def extract(self, file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise ExtractionError(f"File not found: {file_path}")

        return self._read_with_encoding(path.read_bytes())

    def extract_from_bytes(self, content: bytes, filename: str = "") -> str:
        return self._read_with_encoding(content)

    def _read_with_encoding(self, content: bytes) -> str:
        """Try multiple encodings to read text content."""
        for encoding in self.SUPPORTED_ENCODINGS:
            try:
                return content.decode(encoding).strip()
            except (UnicodeDecodeError, LookupError):
                continue
        raise ExtractionError(
            f"Could not decode text file with any of: {self.SUPPORTED_ENCODINGS}"
        )


class ExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


class TextExtractorFactory:
    """
    Factory for selecting the right text extractor based on file type.
    
    Usage:
        extractor = TextExtractorFactory.get("application/pdf")
        text = extractor.extract("/path/to/file.pdf")
    """

    _extractors: dict[str, type[BaseExtractor]] = {
        "application/pdf": PDFExtractor,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor,
        "text/plain": TextExtractor,
        "text/markdown": TextExtractor,
        "text/x-markdown": TextExtractor,
    }

    # Fallback by extension if MIME type is missing
    _extension_map: dict[str, type[BaseExtractor]] = {
        ".pdf": PDFExtractor,
        ".docx": DOCXExtractor,
        ".txt": TextExtractor,
        ".md": TextExtractor,
        ".markdown": TextExtractor,
    }

    @classmethod
    def get(cls, mime_type: str | None = None, filename: str | None = None) -> BaseExtractor:
        """
        Get an extractor for the given MIME type or filename.
        Raises ExtractionError if no matching extractor is found.
        """
        if mime_type and mime_type.lower() in cls._extractors:
            return cls._extractors[mime_type.lower()]()

        if filename:
            ext = Path(filename).suffix.lower()
            if ext in cls._extension_map:
                return cls._extension_map[ext]()

        raise ExtractionError(
            f"No extractor available for MIME type '{mime_type}' / file '{filename}'. "
            f"Supported types: PDF, DOCX, TXT, Markdown"
        )

    @classmethod
    def supported_types(cls) -> list[str]:
        """Return list of supported MIME types."""
        return list(cls._extractors.keys())

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Return list of supported file extensions."""
        return list(cls._extension_map.keys())
